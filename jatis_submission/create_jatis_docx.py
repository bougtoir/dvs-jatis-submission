#!/usr/bin/env python3
"""
Generate JATIS manuscript as docx with inline figures.

Title: "Solving the noise inverse problem in dynamic vision sensors
        for faint astronomical object detection"

Target: Journal of Astronomical Telescopes, Instruments, and Systems (SPIE)

SPIE formatting:
- Single column, Times Roman 12pt
- Numbered references (superscript in text) — here we use [N] for docx clarity
- Single-paragraph abstract (≤200 words)
- Figures inline with captions below
- Section numbering: 1, 1.1, 1.2, etc.
- Equations as Word equation objects (OMML)
"""

import re
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR


# =========================================================
# OMML equation helpers
# =========================================================

def _mr(parent, text, italic=True, bold=False):
    """Create a math run (m:r) with text."""
    r = etree.SubElement(parent, qn('m:r'))
    if not italic or bold:
        rPr = etree.SubElement(r, qn('m:rPr'))
        if not italic:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'p')
        if bold:
            sty = etree.SubElement(rPr, qn('m:sty'))
            sty.set(qn('m:val'), 'b')
    t = etree.SubElement(r, qn('m:t'))
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _sub(parent, base, sub):
    """Create subscript: base_sub."""
    el = etree.SubElement(parent, qn('m:sSub'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sub'))
    _mr(s, sub)
    return el


def _sup(parent, base, sup):
    """Create superscript: base^sup."""
    el = etree.SubElement(parent, qn('m:sSup'))
    e = etree.SubElement(el, qn('m:e'))
    _mr(e, base)
    s = etree.SubElement(el, qn('m:sup'))
    _mr(s, sup)
    return el


def _frac(parent, num_builder, den_builder):
    """Create fraction with builder functions for numerator and denominator."""
    f = etree.SubElement(parent, qn('m:f'))
    num = etree.SubElement(f, qn('m:num'))
    num_builder(num)
    den = etree.SubElement(f, qn('m:den'))
    den_builder(den)
    return f


def _delim(parent, content_builder, left='(', right=')'):
    """Create delimiter (parentheses, brackets)."""
    d = etree.SubElement(parent, qn('m:d'))
    dPr = etree.SubElement(d, qn('m:dPr'))
    begChr = etree.SubElement(dPr, qn('m:begChr'))
    begChr.set(qn('m:val'), left)
    endChr = etree.SubElement(dPr, qn('m:endChr'))
    endChr.set(qn('m:val'), right)
    e = etree.SubElement(d, qn('m:e'))
    content_builder(e)
    return d


def _bar(parent, text):
    """Create accent bar (hat/overline): x̂ or x̄."""
    acc = etree.SubElement(parent, qn('m:acc'))
    accPr = etree.SubElement(acc, qn('m:accPr'))
    chrEl = etree.SubElement(accPr, qn('m:chr'))
    chrEl.set(qn('m:val'), '\u0302')  # combining circumflex (hat)
    e = etree.SubElement(acc, qn('m:e'))
    _mr(e, text)
    return acc


def _hat(parent, text):
    """Create hat accent: x̂."""
    return _bar(parent, text)


def _func(parent, name, arg_builder):
    """Create function application: name(args)."""
    _mr(parent, name, italic=False)
    _delim(parent, arg_builder)


def add_display_equation(doc, builder_func, eq_num=None):
    """Add a display equation (centered) with optional number."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    omathpara = etree.SubElement(p._element, qn('m:oMathPara'))
    omath = etree.SubElement(omathpara, qn('m:oMath'))
    builder_func(omath)
    if eq_num:
        run = p.add_run(f'    ({eq_num})')
        run.font.size = Pt(12)
    return p


def add_inline_math(p, builder_func):
    """Add inline equation within a paragraph."""
    omath = etree.SubElement(p._element, qn('m:oMath'))
    builder_func(omath)
    return omath


# =========================================================
# Equation definitions for this manuscript
# =========================================================

def eq_a5_model(omath):
    """Eq (1): λ_noise(T, I_bg) = I_dark,ref · exp(α·ΔT) · (1 + β·I_bg)"""
    _sub(omath, '\u03bb', 'noise')
    def _args(e):
        _mr(e, 'T, ')
        _sub(e, 'I', 'bg')
    _delim(omath, _args)
    _mr(omath, ' = ')
    _sub(omath, 'I', 'dark,ref')
    _mr(omath, ' \u22c5 exp')
    def _exp_arg(e):
        _mr(e, '\u03b1 \u22c5 \u0394T')
    _delim(omath, _exp_arg)
    _mr(omath, ' \u22c5 ')
    def _bg(e):
        _mr(e, '1 + \u03b2 \u22c5 ')
        _sub(e, 'I', 'bg')
    _delim(omath, _bg)


def eq_accuracy(omath):
    """Eq (2): α = 1 − ‖ê_noise − e_noise,true‖ / ‖e_noise,true‖"""
    _mr(omath, '\u03b1 = 1 \u2212 ')
    def _num(n):
        _mr(n, '\u2016')
        _hat(n, 'e')
        _sub(n, '', 'noise')
        _mr(n, ' \u2212 ')
        _sub(n, 'e', 'noise,true')
        _mr(n, '\u2016')
    def _den(d):
        _mr(d, '\u2016')
        _sub(d, 'e', 'noise,true')
        _mr(d, '\u2016')
    _frac(omath, _num, _den)


def eq_residual_noise(omath):
    """Eq (3): σ_residual = (1 − α) · σ_original"""
    _sub(omath, '\u03c3', 'residual')
    _mr(omath, ' = ')
    def _factor(e):
        _mr(e, '1 \u2212 \u03b1')
    _delim(omath, _factor)
    _mr(omath, ' \u22c5 ')
    _sub(omath, '\u03c3', 'original')


def eq_snr_improvement(omath):
    """Eq (4): SNR_after / SNR_before = 1 / (1 − α)"""
    def _num(n):
        _sub(n, 'SNR', 'after')
    def _den(d):
        _sub(d, 'SNR', 'before')
    _frac(omath, _num, _den)
    _mr(omath, ' = ')
    def _num2(n):
        _mr(n, '1', italic=False)
    def _den2(d):
        _mr(d, '1 \u2212 \u03b1')
    _frac(omath, _num2, _den2)


def eq_map_estimation(omath):
    """Eq (5): θ̂ = argmax_θ p(e_cal|θ) · p(θ|θ_prior)"""
    _hat(omath, '\u03b8')
    _mr(omath, ' = ')
    _sub(omath, 'argmax', '\u03b8')
    _mr(omath, ' p')
    def _likelihood(e):
        _sub(e, 'e', 'cal')
        _mr(e, ' | \u03b8')
    _delim(omath, _likelihood)
    _mr(omath, ' \u22c5 p')
    def _prior(e):
        _mr(e, '\u03b8 | ')
        _sub(e, '\u03b8', 'prior')
    _delim(omath, _prior)


def eq_nn_output(omath):
    """Eq (6): λ̂_noise = λ_physics · (1 + Δλ_aux) + Δλ_corr"""
    _hat(omath, '\u03bb')
    _mr(omath, '')
    _sub(omath, '', 'noise')
    _mr(omath, ' = ')
    _sub(omath, '\u03bb', 'physics')
    _mr(omath, ' \u22c5 ')
    def _modulation(e):
        _mr(e, '1 + \u0394')
        _sub(e, '\u03bb', 'aux')
    _delim(omath, _modulation)
    _mr(omath, ' + \u0394')
    _sub(omath, '\u03bb', 'corr')


def eq_p_noise(omath):
    """Eq (7): P_noise(e_i) = λ̂_noise(x_i,y_i,t_i) / [λ̂_noise(...) + λ̂_signal(...)]"""
    _sub(omath, 'P', 'noise')
    def _ei(e):
        _sub(e, 'e', 'i')
    _delim(omath, _ei)
    _mr(omath, ' = ')
    def _num(n):
        _hat(n, '\u03bb')
        _sub(n, '', 'noise')
        def _coords(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(n, _coords)
    def _den(d):
        _hat(d, '\u03bb')
        _sub(d, '', 'noise')
        def _c1(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(d, _c1)
        _mr(d, ' + ')
        _hat(d, '\u03bb')
        _sub(d, '', 'signal')
        def _c2(e):
            _sub(e, 'x', 'i')
            _mr(e, ', ')
            _sub(e, 'y', 'i')
            _mr(e, ', ')
            _sub(e, 't', 'i')
        _delim(d, _c2)
    _frac(omath, _num, _den)


def eq_fano(omath):
    """Eq (8): F = Var(N_k) / Mean(N_k)"""
    _mr(omath, 'F', italic=True)
    _mr(omath, ' = ', italic=False)
    def _num(n):
        _mr(n, 'Var', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(n, _nk)
    def _den(d):
        _mr(d, 'Mean', italic=False)
        def _nk(e):
            _sub(e, 'N', 'k')
        _delim(d, _nk)
    _frac(omath, _num, _den)


def eq_detection_limit(omath):
    """Eq (9): Δm ≈ 2.5 log₁₀(1/(1−α))"""
    _mr(omath, '\u0394m \u2248 2.5 ')
    _sub(omath, 'log', '10')
    def _arg(e):
        def _n(n):
            _mr(n, '1', italic=False)
        def _d(d):
            _mr(d, '1 \u2212 \u03b1')
        _frac(e, _n, _d)
    _delim(omath, _arg)


# =========================================================
# Document building helpers
# =========================================================

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, style='Normal', bold=False, italic=False,
                  alignment=None, space_after=None, space_before=None):
    """Add paragraph with optional superscript citation handling."""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)

    # Parse {N} markers for superscript citations
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
    return p


def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert figure inline with caption below."""
    if not img_path.exists():
        add_paragraph(doc, f"[MISSING FIGURE: {img_path.name}]")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(12)
    run = p_img.add_run()
    run.add_picture(str(img_path), width=width)

    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(12)
    # Parse caption for superscript refs
    parts = re.split(r'(\{[^}]+\})', caption)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run_cap = p_cap.add_run(part[1:-1])
            run_cap.font.superscript = True
            run_cap.font.size = Pt(8)
        else:
            run_cap = p_cap.add_run(part)
            run_cap.font.size = Pt(9)
    return p_cap


def add_table(doc, headers, data, caption=None):
    """Add a table with headers and data rows."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val
            for p in table.rows[row_idx + 1].cells[col_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if caption:
        add_paragraph(doc, caption, italic=True, space_after=12,
                      space_before=6)
    return table


def build_manuscript():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Page setup (Letter size per SPIE)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.25)
    section.left_margin = Inches(0.875)
    section.right_margin = Inches(0.875)

    # =========================================================
    # Title
    # =========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Solving the noise inverse problem in dynamic vision sensors\n'
        'for faint astronomical object detection'
    )
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('[Author names and affiliations to be added]')
    run.font.size = Pt(12)
    run.italic = True

    # =========================================================
    # Abstract (single paragraph, ≤200 words for SPIE)
    # =========================================================
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) provide microsecond temporal resolution '
        'and >120 dB dynamic range, properties well suited to space situational '
        'awareness (SSA). Under low-light conditions, however, background '
        'activity (BA) noise overwhelms faint astronomical signals. We address '
        'this by formulating DVS denoising as a noise inverse problem: the '
        'noise generation mechanism is described by a circuit-level parametric '
        'model (A5, five parameters per pixel) and solved inversely to '
        'reconstruct and subtract noise, leaving signal in the residual. '
        'Our algorithm, Physics-Informed DeepClean for DVS (PI-DC-DVS), '
        'combines this model with Bayesian inference. On 20 recordings from '
        'the Event-Based Space Situational Awareness (EBSSA) dataset, a '
        'Fano-factor classifier derived from this framework reaches '
        'ROC-AUC = 0.866, compared with 0.534 for conventional temporal '
        'filtering. A proof-of-concept run removes 90.3% of noise events '
        'while recovering satellite trajectories. We also describe a six-tier '
        'calibration hierarchy whose highest tier (Cal-6) turns satellite '
        'light trails\u2014ordinarily dismissed as light pollution\u2014into '
        'calibration sources with known trajectories. The noise inverse '
        'approach, long established in gravitational-wave data analysis, '
        'proves well suited to event-based astronomical observation and '
        'offers a potential 2\u20134 magnitude gain in detection limit.'
    ))

    add_paragraph(doc, (
        'Keywords: dynamic vision sensor, event camera, noise inverse problem, '
        'space situational awareness, calibration, stochastic resonance'
    ), italic=True, space_after=12)

    # =========================================================
    # 1. Introduction
    # =========================================================
    add_heading(doc, '1. Introduction', level=1)
    add_paragraph(doc, (
        'Classical image denoising is a signal inverse problem: given an '
        'observation y = h * s + n, estimate s. This paper takes the '
        'complementary view. We model the noise generation mechanism as a '
        'physical forward process and invert it to reconstruct and subtract '
        'noise, so that only signal remains in the residual.'
    ))
    add_paragraph(doc, (
        'Dynamic Vision Sensors (DVS) are neuromorphic image sensors whose '
        'pixels independently fire events whenever the logarithmic intensity '
        'change crosses a threshold.{1,2} The resulting data stream has '
        'microsecond time resolution, >120 dB dynamic range, and sparse '
        'bandwidth\u2014properties that favour space situational awareness '
        '(SSA){3,4,5} and fast optical astronomy.{6} At low light levels, '
        'however, shot-noise-driven background activity (BA) dominates the '
        'event stream and buries faint astronomical signals.{7,8}'
    ))
    add_paragraph(doc, (
        'Our premise is straightforward: if the noise can be modelled with '
        'high fidelity, subtracting it yields a structural gain in SNR\u2014one '
        'that scales with model accuracy rather than integration time.'
    ), italic=True)
    add_paragraph(doc, (
        'Gravitational-wave (GW) astronomy provides a proven precedent. '
        'DeepClean{9} and successors{10,11,12} regress non-stationary detector '
        'noise from auxiliary witness channels, and iDQ{13} assigns per-event '
        'noise probabilities in real time. For photon-counting data, '
        'D\u00b3PO{14} jointly decomposes signal and noise in a Bayesian framework. '
        'No comparable pipeline has been developed for DVS astronomical '
        'observation.'
    ))
    add_paragraph(doc, (
        'This paper makes five contributions: (i) we formulate DVS denoising '
        'as a noise inverse problem and derive the SNR improvement scaling '
        '(Sec. 3); (ii) we describe PI-DC-DVS, an algorithm that couples a '
        'circuit-level noise model with auxiliary-channel regression (Sec. 3); '
        '(iii) we introduce a six-tier calibration framework whose top tier '
        '(Cal-6) exploits satellite light trails as calibration sources '
        '(Sec. 4); (iv) we evaluate the approach on 20 EBSSA recordings '
        'against two baselines (Sec. 5); and (v) we demonstrate 90.3% noise '
        'removal with satellite trajectory recovery on a representative '
        'recording (Sec. 6).'
    ))

    # =========================================================
    # 2. Background
    # =========================================================
    add_heading(doc, '2. Background', level=1)

    add_heading(doc, '2.1. DVS noise physics', level=2)
    add_paragraph(doc, (
        'The circuit-level physics of DVS noise has been characterised in '
        'detail. Gra\u00e7a and Delbruck{7} established that photon shot noise '
        'sets a fundamental lower bound on background activity rates. '
        'McReynolds et al.{8} showed that shot-noise events come in '
        'alternating ON/OFF polarity pairs. The SciDVS sensor{15} reaches '
        '1.7% temporal contrast sensitivity at 0.7 lux. Of particular '
        'relevance here, Gra\u00e7a and Delbruck{16} developed a large-signal '
        'differential-equation pixel model with first-passage-time event '
        'generation that is >1000\u00d7 faster than SPICE while retaining '
        'physical accuracy. This model provides the forward model F(\u03b8) '
        'for our noise inverse formulation.'
    ))
    add_paragraph(doc, (
        'The parametric noise rate model (hereafter A5){16} takes the form:'
    ))
    add_display_equation(doc, eq_a5_model, eq_num='1')
    add_paragraph(doc, (
        'Here I_dark,ref is the dark-current reference rate, '
        '\u03b1 \u2248 0.06\u20130.08 K\u207b\u00b9 the temperature coefficient, \u0394T the '
        'temperature departure from a reference point, and \u03b2 the illuminance '
        'sensitivity. The five per-pixel parameters '
        '(I_dark,ref, \u03b1, \u03b2, \u03b8_ON, \u03b8_OFF) are determined during '
        'offline calibration.'
    ))

    add_heading(doc, '2.2. DVS denoising methods', level=2)
    add_paragraph(doc, (
        'DVS denoising methods range from spatio-temporal neighbourhood '
        'filters{17,18} through probabilistic approaches (Event Probability '
        'Mask{19}) and deep networks (WedNet,{20} ASTEDNet{21}) to joint '
        'motion\u2013noise estimation. The closest prior work is that of Shiba '
        'et al.,{22} who estimate motion and noise simultaneously within an '
        'extended Contrast Maximisation framework. Their noise model remains '
        'phenomenological, however, and incorporates neither circuit-level '
        'physics nor auxiliary sensor channels.'
    ))

    add_heading(doc, '2.3. DVS astronomical applications', level=2)
    add_paragraph(doc, (
        'DVS work in astronomy has focused on SSA. Afshar et al.{3} '
        'published the first event-based space-observation dataset (EBSSA: '
        '236 recordings, 572 labelled objects). Ralph et al.{4} demonstrated '
        'unsupervised real-time tracking, and Hoang{6} assessed neuromorphic '
        'cameras for atmospheric Cherenkov telescopes. Faint-object detection '
        'in the noise-dominated regime remains unexplored.'
    ))

    # Fig 2: Gap map
    add_figure(doc, FIG_DIR / 'fig2_gapmap.png',
               'Fig. 1. Gap map showing the four surveyed domains (A: DVS noise '
               'physics, B: DVS denoising, C: DVS astronomical applications, '
               'D: noise inverse problem methods) and the identified research gaps '
               'at their intersections.',
               width=Inches(5.0))

    add_heading(doc, '2.4. Noise inverse problem in other fields', level=2)
    add_paragraph(doc, (
        'In gravitational-wave physics, DeepClean{9} regresses '
        'non-stationary detector noise from auxiliary witness channels, '
        'achieving order-of-magnitude noise reduction at LIGO. In the DVS '
        'domain, Noise2Image{23} recovers static scene content from noise '
        'alone by exploiting the illuminance dependence of noise rates\u2014'
        'direct evidence that DVS noise carries usable information. '
        'D\u00b3PO{14} performs joint Bayesian signal\u2013noise decomposition for '
        'photon-counting data within the information field theory formalism.'
    ))

    # =========================================================
    # 3. Methods
    # =========================================================
    add_heading(doc, '3. Methods: PI-DC-DVS algorithm', level=1)

    add_heading(doc, '3.1. Fundamental principle', level=2)
    add_paragraph(doc, (
        'We define a noise model accuracy \u03b1:'
    ))
    add_display_equation(doc, eq_accuracy, eq_num='2')
    add_paragraph(doc, (
        'After subtraction, the residual noise level is:'
    ))
    add_display_equation(doc, eq_residual_noise, eq_num='3')
    add_paragraph(doc, (
        'so that the SNR improves by:'
    ))
    add_display_equation(doc, eq_snr_improvement, eq_num='4')
    add_paragraph(doc, (
        'At \u03b1 = 0.9 the gain is 10\u00d7; at \u03b1 = 0.99, 100\u00d7. Because the '
        'gain depends only on model accuracy, the priority is to build a good '
        'noise model rather than to match signal templates.'
    ))

    add_heading(doc, '3.2. Algorithm overview', level=2)
    add_paragraph(doc, (
        'The PI-DC-DVS pipeline has four phases (Fig. 2):'
    ))

    # Fig 1: Pipeline
    add_figure(doc, FIG_DIR / 'fig1_pipeline.png',
               'Fig. 2. System architecture of the PI-DC-DVS noise inverse problem '
               'pipeline. Four stages: (1) noise forward model construction using the '
               'A5 pixel model and auxiliary channels; (2) Bayesian inverse problem '
               'solution with physics-informed neural network; (3) residual event stream '
               'generation via probabilistic thinning; (4) astronomical calibration and '
               'verification including Cal-6 satellite trail calibration.',
               width=Inches(6.0))

    add_paragraph(doc, (
        'Phase 1 \u2014 Offline calibration. Dark frames (lens cap), flat-field '
        'frames (integrating sphere), and thermal-sweep recordings '
        '(\u0394T = \u00b15\u00b0C) are acquired before observation. Per-pixel parameters '
        'of the A5 forward model{16} are estimated via MAP:'
    ))
    add_display_equation(doc, eq_map_estimation, eq_num='5')
    add_paragraph(doc, (
        'Phase 2 \u2014 Online inference. A three-layer physics-informed neural '
        'network predicts per-pixel noise rates in real time: (a) a physics '
        'layer with fixed weights encoding the A5 baseline, (b) an '
        'auxiliary-channel coupling layer (MLP, 64-32-1 units) that learns '
        'non-stationary corrections, and (c) a spatio-temporal correlation '
        'layer (3\u00d73 Conv2D) that captures inter-pixel dependencies. '
        'The output is:'
    ))
    add_display_equation(doc, eq_nn_output, eq_num='6')
    add_paragraph(doc, (
        'Each event is then assigned a noise probability in the style of '
        'iDQ:{13}'
    ))
    add_display_equation(doc, eq_p_noise, eq_num='7')
    add_paragraph(doc, (
        'Phase 3 \u2014 Residual generation. In soft subtraction, each event '
        'receives a weight w_i = 1 \u2212 P_noise(e_i) and is retained if w_i '
        'exceeds a threshold. In hard subtraction, events with '
        'P_noise(e_i) < \u03c4 are kept directly.'
    ))
    add_paragraph(doc, (
        'Phase 4 \u2014 Adaptive updates. Residual Poisson statistics are '
        'monitored and network weights updated online via Kalman-filter '
        'drift correction.'
    ))

    add_heading(doc, '3.3. Simplified implementation for EBSSA', level=2)
    add_paragraph(doc, (
        'Fano filter (proposed baseline). The Fano factor\u2014the '
        'variance-to-mean ratio of event counts across temporal bins\u2014'
        'separates noise from signal: Poisson-consistent pixels (F \u2248 1) are '
        'noise-dominated, while F \u226b 1 flags bursty, signal-bearing '
        'activity. Formally:'
    ))
    add_display_equation(doc, eq_fano, eq_num='8')
    add_paragraph(doc, (
        'with N_k the event count in bin k. Noise-dominated pixels '
        '(F \u2264 2) supply the empirical noise-rate estimate; events with '
        'P_noise > \u03c4 (\u03c4 = 0.5) are classified as noise.'
    ))
    add_paragraph(doc, (
        'Simplified PI-DC-DVS neural network. A three-layer network\u2014'
        'physics layer, temporal modulation layer (standing in for absent '
        'auxiliary channels), and spatio-temporal correlation layer\u2014is '
        'trained self-supervised on noise-dominated pixels with a Poisson '
        'negative log-likelihood loss.'
    ))

    # =========================================================
    # 4. Calibration framework
    # =========================================================
    add_heading(doc, '4. Calibration framework', level=1)

    add_heading(doc, '4.1. Six-tier calibration dataset', level=2)
    add_paragraph(doc, (
        'Because DVS produce asynchronous event streams rather than frames, '
        'dedicated calibration procedures are needed. We define a six-tier '
        'hierarchy (Table 1):'
    ))

    add_table(doc,
              ['Tier', 'Condition', 'Purpose', 'Pass criterion'],
              [
                  ['Cal-1', 'Dark (lens cap)', 'Pure noise reference', '\u03c7\u00b2/dof < 1.5'],
                  ['Cal-2', 'Thermal sweep', 'Temperature dependence', 'Residual < 10%'],
                  ['Cal-3', 'Flat-field', 'Shot noise statistics', '\u03b1_flat > 0.9'],
                  ['Cal-4', 'Dynamic patterns', 'Injection-recovery', 'AUC > 0.95'],
                  ['Cal-5', 'Simulated astro.', 'End-to-end pipeline', '\u0394m > 2 mag'],
                  ['Cal-6', 'Satellite trails', 'In-operation verification', 'Det. rate > 95%'],
              ],
              caption='Table 1. Six-tier calibration framework for DVS noise model validation.')

    add_heading(doc, '4.2. Cal-6: Satellite trail calibration', level=2)
    add_paragraph(doc, (
        'Satellite light trails are ordinarily treated as light pollution. '
        'We argue they are, in fact, useful calibration sources. Artificial '
        'satellites follow precisely predicted orbits: Two-Line Element (TLE) '
        'data give position and transit time to sub-arcsecond, '
        'sub-millisecond accuracy, making every transit a natural '
        'injection-recovery test under real sky conditions.'
    ))
    add_paragraph(doc, (
        'DVS sensors are well matched to this application: (1) Starlink and '
        'other constellations supply dozens of transits per night at any '
        'site; (2) the trails are recorded under realistic atmospheric and '
        'background conditions; (3) the >120 dB dynamic range of DVS avoids '
        'the saturation that plagues CCDs; and (4) no extra hardware or '
        'observing time is required.'
    ))
    add_paragraph(doc, (
        'The procedure is: predict the transit via TLE + SGP4; extract the '
        'event stream within the transit window; run the noise subtraction '
        'pipeline; compare the detected trail against the predicted '
        'trajectory; and quantify detection rate, positional accuracy, '
        'and timing residuals.'
    ))

    # =========================================================
    # 5. Systematic evaluation
    # =========================================================
    add_heading(doc, '5. Systematic evaluation', level=1)

    add_heading(doc, '5.1. Dataset', level=2)
    add_paragraph(doc, (
        'We draw on the EBSSA dataset:{3} 236 recordings from DAVIS240C '
        'sensors observing satellites and stars, with 572 labelled objects. '
        'We select 20 recordings spanning both sensor formats '
        '(180\u00d7240 and 240\u00d7304 pixels).'
    ))

    add_heading(doc, '5.2. Evaluated methods', level=2)
    add_paragraph(doc, (
        'We compare three methods: (1) the Fano filter (proposed), which '
        'uses the Fano factor [Eq. (8)] as a Poisson discriminant; (2) a '
        'simplified PI-DC-DVS neural network, self-supervised on '
        'noise-dominated pixels; and (3) the temporal neighbourhood filter '
        'of Delbruck{17} (baseline), which retains an event only if enough '
        'neighbours fall within a fixed spatio-temporal window.'
    ))

    add_heading(doc, '5.3. Results', level=2)

    add_table(doc,
              ['Method', 'NRR', 'SPR', 'F1', 'AUC'],
              [
                  ['Temporal filter', '0.852 \u00b1 0.044', '0.216 \u00b1 0.157',
                   '0.253 \u00b1 0.176', '0.534 \u00b1 0.083'],
                  ['PI-DC-DVS NN', '0.171 \u00b1 0.342', '0.841 \u00b1 0.342',
                   '0.488 \u00b1 0.453', '0.546 \u00b1 0.218'],
                  ['Fano filter', '0.713 \u00b1 0.232', '0.939 \u00b1 0.056',
                   '0.697 \u00b1 0.339', '0.866 \u00b1 0.107'],
              ],
              caption='Table 2. Systematic evaluation results (mean \u00b1 std) across '
                      '20 EBSSA recordings. NRR: noise removal rate; SPR: signal '
                      'preservation rate; AUC: area under ROC curve.')

    add_paragraph(doc, (
        'The Fano filter gives the best trade-off between noise removal and '
        'signal preservation (Fig. 3), reaching AUC = 0.866 against 0.534 for '
        'the temporal filter and 0.546 for the simplified neural network. The '
        'temporal filter does remove the most noise (NRR = 85.2%) but at the '
        'expense of most signal events as well (SPR = 21.6%), making it '
        'unsuitable for faint-object work.'
    ))

    # Fig 3: Evaluation
    add_figure(doc, FIG_DIR / 'fig3_evaluation.png',
               'Fig. 3. Systematic evaluation of three denoising methods on 20 EBSSA '
               'recordings. Four-panel boxplot showing (a) Noise Removal Rate, '
               '(b) Signal Preservation Rate, (c) F1 Score, and (d) ROC-AUC. '
               'Diamond markers indicate means. The Fano filter achieves the best '
               'overall balance (AUC = 0.866).',
               width=Inches(5.5))

    add_paragraph(doc, (
        'Recording-by-recording results (Fig. 4) confirm that the Fano '
        'filter outperforms the temporal filter across a range of noise '
        'rates, background levels, and target brightnesses.'
    ))

    # Fig 5: Per-recording
    add_figure(doc, FIG_DIR / 'fig5_per_recording.png',
               'Fig. 4. Per-recording noise removal rate comparison across 20 EBSSA '
               'recordings. The Fano filter (orange) achieves selective noise removal, '
               'while the temporal filter (green) removes events indiscriminately.',
               width=Inches(5.5))

    add_heading(doc, '5.4. A5-based noise rate simulation', level=2)
    add_paragraph(doc, (
        'We use the A5 model [Eq. (1)] to map noise rates and SNR gains '
        'across temperature and illuminance '
        '(T \u2208 [10, 65] \u00b0C, I_bg \u2208 [0.1, 1000] lux). Figure 5 shows '
        'the predicted noise-rate surface, the SNR gain at \u03b1 = 0.9, and '
        'the temperature dependence at fixed illuminance. The model predicts '
        'a mean 5.4\u00d7 SNR improvement (peak 10.0\u00d7) at \u03b1 = 0.9, consistent '
        'with the empirical Fano filter results.'
    ))

    # Fig 4: A5 simulation
    add_figure(doc, FIG_DIR / 'fig4_a5_simulation.png',
               'Fig. 5. A5-based noise rate simulation across the temperature\u2013'
               'illuminance parameter space. (a) Predicted noise event rate '
               '[evt/s/pix]; (b) SNR improvement factor at 90% noise model accuracy; '
               '(c) SNR vs. temperature at fixed illuminance (I_bg = 16.7 lux) '
               'comparing raw, Fano filter, and PI-DC-DVS methods.',
               width=Inches(6.0))

    # =========================================================
    # 6. Proof-of-concept demonstration
    # =========================================================
    add_heading(doc, '6. Proof-of-concept demonstration', level=1)
    add_paragraph(doc, (
        'Starting from 1,800,674 events, probabilistic thinning at '
        '\u03c4 = 0.5 retains 175,261 (9.7%), removing 90.3% of the noise. '
        'Satellite trajectories that are invisible in the raw event '
        'accumulation become clearly visible in the residual (Fig. 6). The '
        'four panels show: (a) raw event accumulation, dominated by BA; '
        '(b) the fitted per-pixel noise-rate map \u03bb_noise(x, y); '
        '(c) the P_noise distribution, bimodal with clean separation at '
        '\u03c4 = 0.5; and (d) the residual event stream with satellite '
        'tracks recovered.'
    ))

    # Fig 6: Demo
    add_figure(doc, FIG_DIR / 'fig6_demo.png',
               'Fig. 6. Proof-of-concept noise inverse problem demonstration on '
               'EBSSA Recording #0. (a) Raw event accumulation (1,800,674 events); '
               '(b) estimated noise rate map \u03bb_noise(x,y) [events/sec]; '
               '(c) per-event noise probability distribution showing bimodal '
               'separation (signal: 9.7%, noise: 90.3%); (d) residual events after '
               '90.3% noise removal (175,261 events) with satellite trajectory '
               'clearly visible.',
               width=Inches(5.5))

    add_paragraph(doc, (
        'The Fano factor map (Fig. 7a) separates noise-dominated pixels '
        '(F \u2248 1, blue) from signal-bearing pixels (F \u226b 1, red). Of '
        '76,800 pixels, 2,294 are flagged as signal candidates, and these '
        'cluster along known satellite tracks. The temporal profile '
        '(Fig. 7b) shows the noise model converging within the first few '
        'time bins. The per-pixel SNR distribution (Fig. 7c) shifts upward '
        'after subtraction, confirming a structural gain.'
    ))

    # Fig 7: SNR
    add_figure(doc, FIG_DIR / 'fig7_snr.png',
               'Fig. 7. Signal-to-noise ratio improvement analysis. (a) Fano factor '
               'spatial map showing noise-dominated (blue, F \u2248 1) vs. signal-'
               'containing pixels (red, F \u226b 1); (b) temporal dynamics of total event '
               'rate vs. noise model prediction; (c) per-pixel SNR distribution '
               'before (raw) and after (residual) noise subtraction.',
               width=Inches(6.0))

    # =========================================================
    # 7. Discussion
    # =========================================================
    add_heading(doc, '7. Discussion', level=1)

    add_heading(doc, '7.1. Effectiveness of the noise inverse problem paradigm', level=2)
    add_paragraph(doc, (
        'The Fano filter (AUC = 0.866) outperforms temporal filtering '
        '(AUC = 0.534) by a wide margin, confirming that physics-based '
        'noise modelling is more productive than pattern-based event '
        'filtering. The Fano factor provides a direct physical '
        'discriminant: noise follows Poisson statistics (F \u2248 1) [Eq. (8)], '
        'while astronomical signals produce bursty event clusters '
        '(F \u226b 1). The parallel with auxiliary-channel methods in '
        'gravitational-wave detectors{9} is instructive: in both cases '
        'the noise has observable statistical structure that permits '
        'separation from signal.'
    ))

    add_heading(doc, '7.2. Role of auxiliary channels', level=2)
    add_paragraph(doc, (
        'Without auxiliary channels, the simplified PI-DC-DVS neural '
        'network barely exceeds chance (AUC = 0.546) and shows large '
        'recording-to-recording variance. This mirrors the LIGO '
        'experience:{9} auxiliary channel information is essential for '
        'accurate noise regression. The Fano filter sidesteps this '
        'requirement by using the physics-based noise rate itself as an '
        'implicit auxiliary channel. Adding real temperature and illuminance '
        'sensors should improve performance further.'
    ))

    add_heading(doc, '7.3. Cal-6: Paradigm inversion of light pollution', level=2)
    add_paragraph(doc, (
        'Cal-6 reverses the usual framing of satellite trails. The growth '
        'of mega-constellations (Starlink, OneWeb) means that dozens of '
        'bright transits cross any observatory field per night, each with '
        'a precisely known trajectory. CCD sensors saturate on these trails, '
        'but DVS record them quantitatively over the full >120 dB range. '
        'At a time when the astronomical community is debating the impact '
        'of satellite constellations on observing,{24} DVS users can turn '
        'these trails into calibration assets.'
    ))

    add_heading(doc, '7.4. Template-free detection via noise residuals', level=2)
    add_paragraph(doc, (
        'High-fidelity noise subtraction enables template-free detection: '
        'any structure that survives in the residual is, by construction, '
        'signal, regardless of morphology. Paired with event-level '
        'shift-and-stack algorithms,{25} this could push the detection of '
        'fast, faint objects (10\u201350 m class near-Earth objects) beyond '
        'the reach of frame-based surveys. From Eq. (4), the '
        'detection-limit gain is:'
    ))
    add_display_equation(doc, eq_detection_limit, eq_num='9')
    add_paragraph(doc, (
        'giving \u0394m > 2.5 mag at \u03b1 = 0.9.'
    ))

    add_heading(doc, '7.5. Limitations and future work', level=2)
    add_paragraph(doc, (
        'The main limitation is the absence of auxiliary channel data in '
        'the EBSSA recordings (neither temperature nor illuminance was '
        'logged). Directions for future work include: '
        '(i) a differentiable A5 pixel model for end-to-end gradient-based '
        'training; '
        '(ii) an auxiliary-channel hardware module (temperature sensor, '
        'photometer) for telescope-mounted DVS; '
        '(iii) on-sky tests with SciDVS{15} on a 0.3\u20130.5 m aperture; '
        '(iv) Cal-6 validation using Starlink transits; and '
        '(v) cross-dataset evaluation on DVSNOISE20.{19}'
    ))

    # =========================================================
    # 8. Conclusions
    # =========================================================
    add_heading(doc, '8. Conclusions', level=1)
    add_paragraph(doc, (
        'This paper has introduced PI-DC-DVS, a physics-informed framework '
        'that formulates DVS denoising as a noise inverse problem. The '
        'principal findings are:'
    ))
    conclusions = [
        '(1) Noise inverse formulation. The SNR gain scales as '
        '1/(1\u2212\u03b1) [Eq. (4)], giving 10\u00d7 at \u03b1 = 0.9 and 100\u00d7 at '
        '\u03b1 = 0.99, with a corresponding 2\u20134 magnitude extension '
        'of the detection limit [Eq. (9)].',

        '(2) Fano-factor classification. The Fano filter reaches '
        'AUC = 0.866 on 20 EBSSA recordings, against 0.534 for temporal '
        'filtering and 0.546 for a simplified neural network.',

        '(3) Proof of concept. Probabilistic thinning removes 90.3% of '
        'noise events while preserving satellite trajectories in the '
        'residual stream.',

        '(4) A5 noise-rate simulation. The parametric model [Eq. (1)] '
        'predicts mean SNR gains of 5.4\u00d7 (peak 10.0\u00d7) at \u03b1 = 0.9, '
        'consistent with the measured Fano filter results.',

        '(5) Six-tier calibration hierarchy. Cal-1 through Cal-6 '
        'progressively validate the noise model; Cal-6 turns satellite '
        'light trails into calibration sources under operational conditions.',
    ]
    for c in conclusions:
        add_paragraph(doc, c, space_after=4)

    # =========================================================
    # Code/Data availability
    # =========================================================
    add_heading(doc, 'Code and data availability', level=1)
    add_paragraph(doc, (
        'The EBSSA dataset is publicly available via the Tonic library.{3} '
        'The implementation code and evaluation scripts are available at '
        'https://github.com/bougtoir/dvs_noise_inverse_problem.'
    ))

    # =========================================================
    # Acknowledgments
    # =========================================================
    add_heading(doc, 'Acknowledgments', level=1)
    add_paragraph(doc, '[To be added.]', italic=True)

    # =========================================================
    # References (SPIE numbered style)
    # =========================================================
    add_heading(doc, 'References', level=1)
    references = [
        # [1] DVS survey
        '[1] G. Gallego, T. Delbr\u00fcck, G. Orchard et al., "Event-based vision: '
        'A survey," IEEE Trans. Pattern Anal. Mach. Intell. 44, 154\u2013180 (2022).',

        # [2] DVS pixel model intro
        '[2] T. Delbruck, C. Li, R. Gra\u00e7a, and B. McReynolds, "Utility and '
        'feasibility of a center surround event camera," in Proc. IEEE ICIP (2022).',

        # [3] EBSSA
        '[3] S. Afshar, A. P. Nicholson, A. van Schaik, and G. Cohen, '
        '"Event-based object detection and tracking for space situational '
        'awareness," preprint arXiv:1911.08730 (2019).',

        # [4] FIESTA
        '[4] N. Ralph, D. Joubert, A. Jolley, S. Afshar, N. Tothill, '
        'A. van Schaik, and G. Cohen, "Real-time event-based unsupervised '
        'feature consolidation and tracking for space situational awareness," '
        'Front. Neurosci. 16, 821157 (2022).',

        # [5] Kaminski / Gedek
        '[5] K. Kami\u0144ski, G. Cohen, T. Delbr\u00fcck, M. \u017bo\u0142nowski, '
        'and M. G\u0119dek, "Observational evaluation of event cameras '
        'performance in optical space surveillance," in Proc. 1st NEO and '
        'Debris Detection Conf., ESA (2019).',

        # [6] Hoang
        '[6] J. Hoang, "Neuromorphic cameras for atmospheric Cherenkov '
        'telescopes and fast optical astronomy: new paradigm, challenges and '
        'opportunities," preprint arXiv:2310.16321 (2023).',

        # [7] Graca 2021 shot noise
        '[7] R. Gra\u00e7a and T. Delbruck, "Unraveling the paradox of intensity-'
        'dependent DVS pixel noise," preprint arXiv:2109.08640 (2021).',

        # [8] McReynolds shot noise pairs
        '[8] B. McReynolds, R. Gra\u00e7a, and T. Delbruck, "Exploiting '
        'alternating DVS shot noise event pair statistics to reduce background '
        'activity rates," preprint arXiv:2304.03494 (2023).',

        # [9] DeepClean
        '[9] G. Vajente, Y. Huang, M. Isi et al., "Machine-learning nonstationary '
        'noise out of gravitational-wave detectors," Phys. Rev. D 101, '
        '042003 (2020).',

        # [10] Dooney
        '[10] T. Dooney, H. Narola, S. Bromuri, R. L. Curier, '
        'C. Van Den Broeck, S. Caudill, and D. S. Tan, "Time-domain '
        'reconstruction of signals and glitches in gravitational wave data '
        'with deep learning," preprint arXiv:2501.18423 (2025).',

        # [11] Wang / WaveFormer
        '[11] H. Wang, Y. Zhou, Z. Cao, Z. Guo, and Z. Ren, "WaveFormer: '
        'transformer-based denoising method for gravitational-wave data," '
        'Mach. Learn.: Sci. Technol. 5, 015046 (2024).',

        # [12] Chatterjee
        '[12] C. Chatterjee and K. Jani, "No glitch in the matrix: robust '
        'reconstruction of gravitational wave signals under noise artifacts," '
        'Astrophys. J. 982, 102 (2025).',

        # [13] iDQ
        '[13] R. Essick, P. Godwin, C. Hanna, L. Blackburn, and '
        'E. Katsavounidis, "iDQ: Statistical inference of non-Gaussian noise '
        'with auxiliary degrees of freedom in gravitational-wave detectors," '
        'Mach. Learn.: Sci. Technol. 2, 015004 (2021).',

        # [14] D3PO
        '[14] M. Selig and T. A. En\u00dflin, "D\u00b3PO\u2014Denoising, deconvolving, and '
        'decomposing photon observations," Astron. Astrophys. 574, A74 (2015).',

        # [15] SciDVS
        '[15] R. Gra\u00e7a, S. Zhou, B. McReynolds, and T. Delbruck, "SciDVS: '
        'A scientific event camera with 1.7% temporal contrast sensitivity '
        'at 0.7 lux," ESSERC 2024, DOI:10.1109/esserc62670.2024.10719521.',

        # [16] Graca 2025 DVS pixel model
        '[16] R. Gra\u00e7a and T. Delbruck, "Towards a physically realistic '
        'computationally efficient DVS pixel model," preprint '
        'arXiv:2505.07386 (2025).',

        # [17] Temporal filter (Delbruck 2008)
        '[17] T. Delbruck, "Frame-free dynamic digital vision," in Proc. Intl. '
        'Symp. on Secure-Life Electronics, 21\u201326 (2008).',

        # [18] Liu 2010
        '[18] S.-C. Liu and T. Delbruck, "Neuromorphic sensory systems," '
        'Curr. Opin. Neurobiol. 20, 288\u2013295 (2010).',

        # [19] EPM / Baldwin
        '[19] R. W. Baldwin, M. Almatrafi, V. Asari, and K. Hirakawa, '
        '"Event probability mask (EPM) and event denoising convolutional '
        'neural network (EDnCNN)," Proc. CVPR (2020).',

        # [20] WedNet
        '[20] H. Fang, J. Wu, Q. Hou, W. Dong, and G. Shi, "Fast window-based '
        'event denoising with spatiotemporal correlation enhancement," '
        'IEEE Trans. Pattern Anal. Mach. Intell. (2024).',

        # [21] ASTEDNet
        '[21] W. Wu, H. Yao, C. Zhai, Z. Dai, and X. Zhu, "Event camera '
        'denoising using asynchronous spatio-temporal event denoising neural '
        'network," ISPRS Archives XLVIII-4-2024 (2024).',

        # [22] Shiba
        '[22] S. Shiba, Y. Aoki, and G. Gallego, "Secrets of event-based '
        'optical flow," in Proc. ECCV (2022).',

        # [23] Noise2Image
        '[23] D. Galor, R. Cao, A. P. Kohli, J. L. Yates, and L. Waller, '
        '"Noise2Image: noise-enabled static scene recovery for event '
        'cameras," Optica 12, 46\u201355 (2025).',

        # [24] Satellite constellation impacts
        '[24] J. McDowell, "The low Earth orbit satellite population and '
        'impacts of the SpaceX Starlink constellation," Astrophys. J. Lett. '
        '892, L36 (2020).',

        # [25] Shift-and-stack
        '[25] S. Stetzler, M. Juri\u0107, P. H. Bernardinelli et al., '
        '"An efficient shift-and-stack algorithm applied to detection '
        'catalogs," Astron. J. 170, 352 (2025).',
    ]
    for ref in references:
        add_paragraph(doc, ref, space_after=3)

    # Save
    out_path = OUT_DIR / 'manuscript_jatis.docx'
    doc.save(str(out_path))
    print(f"Manuscript saved: {out_path}")


if __name__ == '__main__':
    build_manuscript()
